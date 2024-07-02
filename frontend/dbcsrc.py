# import matplotlib
import datetime
import glob
import io
import json
import os
import time
from ast import literal_eval

import docx
import pandas as pd
import requests

# from snapshot_selenium
import snapshot as driver
import streamlit as st
from bs4 import BeautifulSoup
from checkrule import get_lawdtlbyid, get_rulelist_byname
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from pyecharts import options as opts
from pyecharts.charts import Bar, Line
from pyecharts.render import make_snapshot

# from snapshot_phantomjs import snapshot as driver
from streamlit_echarts import st_pyecharts
from utils import df2aggrid, df2echartstable, split_words

pencsrc = "../data/penalty/csrc"
# mapfolder = 'data/temp/citygeo.csv'

BASE_URL = (
    "https://neris.csrc.gov.cn/falvfagui/multipleFindController/solrSearchWrit?pageNo="
)

urldbase = (
    "https://neris.csrc.gov.cn/falvfagui/rdqsHeader/lawWritInfo?navbarId=1&lawWritId="
)

jspath = "{}/".format(os.path.dirname(os.path.abspath("data/map/echarts.min.js")))


# @st.cache(allow_output_mutation=True)
def get_csvdf(penfolder, beginwith):
    files2 = glob.glob(penfolder + "**/" + beginwith + "*.csv", recursive=True)
    dflist = []
    # filelist = []
    for filepath in files2:
        pendf = pd.read_csv(filepath)
        dflist.append(pendf)
        # filelist.append(filename)
    if len(dflist) > 0:
        df = pd.concat(dflist)
        df.reset_index(drop=True, inplace=True)
    else:
        df = pd.DataFrame()
    return df


# @st.cache(allow_output_mutation=True)
def get_csrcdetail():
    pendf = get_csvdf(pencsrc, "sdresult")
    # format date
    pendf["发文日期"] = pd.to_datetime(pendf["发文日期"]).dt.date
    # sort by date
    pendf = pendf.sort_values(by="发文日期", ascending=False)
    # fillna
    pendf = pendf.fillna("")
    return pendf


# @st.cache(allow_output_mutation=True)
def get_csrcsum():
    pendf = get_csvdf(pencsrc, "sumevent")
    # format date
    pendf["date1"] = pd.to_datetime(pendf["date"], errors="coerce")
    pendf.loc[pendf["date1"].isnull(), "date1"] = pd.to_datetime(
        pendf.loc[pendf["date1"].isnull(), "date"], unit="ms", errors="coerce"
    ).dt.date

    # Check for remaining nulls
    null_rows = pendf[pendf["date1"].isnull()]
    if not null_rows.empty:
        print("Rows where date conversion failed:")
        print(null_rows)

    return pendf


# get lawdetail
# @st.cache(allow_output_mutation=True)
def get_lawdetail():
    lawdf = get_csvdf(pencsrc, "lawdf")
    # format date
    lawdf["发文日期"] = pd.to_datetime(lawdf["发文日期"]).dt.date
    # fillna
    lawdf = lawdf.fillna("")
    return lawdf


# get peopledetail
# @st.cache(allow_output_mutation=True)
def get_peopledetail():
    peopledf = get_csvdf(pencsrc, "peopledf")
    # format date
    peopledf["发文日期"] = pd.to_datetime(peopledf["发文日期"]).dt.date
    # fillna
    peopledf = peopledf.fillna("")
    return peopledf


# search by filename, date, org, case, type
def searchcsrc(df, filename, start_date, end_date, org, case, type):
    col = ["文件名称", "发文日期", "发文单位", "案情经过", "文书类型", "id"]
    # convert date to datetime
    # df['发文日期'] = pd.to_datetime(df['发文日期']).dt.date
    # split words
    if filename != "":
        filename = split_words(filename)
    if org != "":
        org = split_words(org)
    if case != "":
        case = split_words(case)

    searchdf = df[
        (df["文件名称"].str.contains(filename))
        & (df["发文日期"] >= start_date)
        & (df["发文日期"] <= end_date)
        & (df["发文单位"].str.contains(org))
        & (df["案情经过"].str.contains(case))
        & (df["文书类型"].isin(type))
    ][col]
    # get summary
    # searchdf1['案情经过'] = searchdf1['案情经过'].apply(get_summary)
    # searchdf1['案情经过'] = searchdf1['案情经过'].apply(lambda x: x[:100] + '...')
    # sort by date desc
    searchdf.sort_values(by=["发文日期"], ascending=False, inplace=True)
    # reset index
    searchdf.reset_index(drop=True, inplace=True)

    return searchdf


# search law by filename_text,start_date,end_date , org_text,law_text,article_text,  type_text
def searchlaw(
    df, filename_text, start_date, end_date, org_text, law_text, article_text, type_text
):
    col = ["文件名称", "发文日期", "文书类型", "发文单位", "法律法规", "条文", "id"]
    # convert date to datetime
    # df['发文日期'] = pd.to_datetime(df['发文日期']).dt.date
    # split words
    if filename_text != "":
        filename_text = split_words(filename_text)
    if org_text != "":
        org_text = split_words(org_text)
    if article_text != "":
        article_text = split_words(article_text)

    searchdf = df[
        (df["文件名称"].str.contains(filename_text))
        & (df["发文日期"] >= start_date)
        & (df["发文日期"] <= end_date)
        & (df["发文单位"].str.contains(org_text))
        & (df["法律法规"].isin(law_text))
        & (df["条文"].str.contains(article_text))
        & (df["文书类型"].isin(type_text))
    ][col]
    # sort by date desc
    searchdf.sort_values(by=["发文日期"], ascending=False, inplace=True)
    # reset index
    searchdf.reset_index(drop=True, inplace=True)
    return searchdf


# search people by filename_text,start_date,end_date , org_text,people_type_text, people_name_text, people_position_text, penalty_type_text, penalty_result_text, type_text)
def searchpeople(
    df,
    filename_text,
    start_date,
    end_date,
    org_text,
    people_type_text,
    people_name_text,
    people_position_text,
    penalty_type_text,
    penalty_result_text,
    type_text,
):
    col = [
        "文件名称",
        "发文日期",
        "文书类型",
        "发文单位",
        "当事人类型",
        "当事人名称",
        "当事人身份",
        "违规类型",
        "处罚结果",
        "id",
    ]
    # convert date to datetime
    # df['发文日期'] = pd.to_datetime(df['发文日期']).dt.date
    # split words
    if filename_text != "":
        filename_text = split_words(filename_text)
    if org_text != "":
        org_text = split_words(org_text)
    if people_name_text != "":
        people_name_text = split_words(people_name_text)
    if penalty_result_text != "":
        penalty_result_text = split_words(penalty_result_text)

    searchdf = df[
        (df["文件名称"].str.contains(filename_text))
        & (df["发文日期"] >= start_date)
        & (df["发文日期"] <= end_date)
        & (df["发文单位"].str.contains(org_text))
        & (df["当事人类型"].isin(people_type_text))
        & (df["当事人名称"].str.contains(people_name_text))
        & (df["当事人身份"].isin(people_position_text))
        & (df["违规类型"].isin(penalty_type_text))
        & (df["处罚结果"].str.contains(penalty_result_text))
        & (df["文书类型"].isin(type_text))
    ][col]
    # sort by date desc
    searchdf.sort_values(by=["发文日期"], ascending=False, inplace=True)
    # reset index
    searchdf.reset_index(drop=True, inplace=True)
    return searchdf


# convert eventdf to lawdf
def generate_lawdf(eventdf):
    law1 = eventdf[
        [
            "id",
            "文件名称",
            "文号",
            "发文日期",
            "文书类型",
            "发文单位",
            "原文链接",
            "处理依据",
        ]
    ]

    law1["处理依据"] = law1["处理依据"].apply(literal_eval)

    law2 = law1.explode("处理依据")

    law3 = law2["处理依据"].apply(pd.Series)

    law4 = pd.concat([law2, law3], axis=1)

    law5 = law4.explode("条文")

    law6 = law5.drop(["处理依据"], axis=1)

    lawdf = law6[
        [
            "id",
            "文件名称",
            "文号",
            "发文日期",
            "文书类型",
            "发文单位",
            "原文链接",
            "法律法规",
            "条文",
        ]
    ]

    # reset index
    lawdf.reset_index(drop=True, inplace=True)
    savedf(lawdf, "lawdf")
    return lawdf


# convert eventdf to peopledf
def generate_peopledf(eventdf):
    peopledf = eventdf[
        [
            "id",
            "文件名称",
            "文号",
            "发文日期",
            "文书类型",
            "发文单位",
            "原文链接",
            "当事人信息",
        ]
    ]

    peopledf["当事人信息"] = peopledf["当事人信息"].apply(literal_eval)

    peopledf2 = peopledf.explode("当事人信息")

    peoplesp1 = peopledf2["当事人信息"].apply(pd.Series)

    peopledf3 = pd.concat([peopledf2, peoplesp1], axis=1)

    peopledf4 = peopledf3[
        [
            "id",
            "文件名称",
            "文号",
            "发文日期",
            "文书类型",
            "发文单位",
            "原文链接",
            "当事人类型",
            "当事人名称",
            "当事人身份",
            "违规类型",
            "处罚结果",
        ]
    ]

    # reset index
    peopledf4.reset_index(drop=True, inplace=True)
    savedf(peopledf4, "peopledf")
    return peopledf4


def savedf(df, basename):
    savename = basename + ".csv"
    savepath = os.path.join(pencsrc, savename)
    df.to_csv(savepath)


# count the number of df by month
def count_by_month(df):
    df_month = df.copy()
    # count by month
    df_month["month"] = df_month["发文日期"].apply(lambda x: x.strftime("%Y-%m"))
    df_month_count = (
        df_month.groupby(["month"])["id"].nunique().reset_index(name="count")
    )
    return df_month_count


# sum amount column of df by month
def sum_amount_by_month(df):
    df["发文日期"] = pd.to_datetime(df["发文日期"]).dt.date
    # df=df[df['发文日期']>=pd.to_datetime('2020-01-01')]
    df["month"] = df["发文日期"].apply(lambda x: x.strftime("%Y-%m"))
    df["amt"] = df["处罚结果"].str.extract("(.*)万元")
    df["amt"].fillna(0, inplace=True)
    df["amt"] = df["amt"].astype(float)
    df["amt"] = df["amt"] * 10000
    df_month_sum = df.groupby(["month"])["amt"].sum().reset_index(name="sum")
    df_sigle_penalty = df.groupby("id")["amt"].sum().reset_index()
    return df_month_sum, df_sigle_penalty


# display searchdf in plotly
def display_dfmonth(search_df):
    # get search_df id list
    search_df_id = search_df["id"].tolist()
    # get people detail
    peopledf = get_peopledetail()
    # search people detail by selected_rows_id
    selected_peopledetail = peopledf[peopledf["id"].isin(search_df_id)]
    # get lawdf
    lawdf = get_lawdetail()
    # search lawdetail by selected_rows_id
    selected_lawdetail = lawdf[lawdf["id"].isin(search_df_id)]
    # get eventdf count by month
    df_month = count_by_month(search_df)
    # get eventdf sum amount by month
    df_sum, df_sigle_penalty = sum_amount_by_month(selected_peopledetail)
    # display checkbox to show/hide graph1
    # showgraph1 = st.sidebar.checkbox("案例数量和金额统计", key="showgraph1")
    showgraph1 = True
    if showgraph1:
        # fig = go.Figure()
        # trace1 = go.Bar(x=df_month['month'], y=df_month['count'], name='数量统计')
        # trace2 = go.Scatter(x=df_sum['month'], y=df_sum['sum'], name='金额统计')
        # fig = make_subplots(specs=[[{"secondary_y": True}]])
        # fig.add_trace(trace1)
        # fig.add_trace(trace2, secondary_y=True)
        # fig.update_layout(height=600,width=800,title_text="案例数量和金额统计")
        # st.plotly_chart(fig)
        x_data = df_month["month"].tolist()
        y_data = df_month["count"].tolist()
        sum_data = df_sum["sum"].tolist()

        # bar = (
        #     Bar()
        #     .add_xaxis(xaxis_data=x_data)
        #     .add_yaxis(series_name="数量", y_axis=y_data, yaxis_index=0)
        #     .set_global_opts(title_opts=opts.TitleOpts(title="案例数量统计"))
        # )
        # line = (
        #     Line()
        #     .add_xaxis(x_data)
        #     .add_yaxis("金额", sum_data, label_opts=opts.LabelOpts(is_show=False))
        #     .set_global_opts(
        #         title_opts=opts.TitleOpts(title="案例金额统计"),
        #         # legend_opts=opts.LegendOpts(pos_top="48%"),
        #     )
        # )
        # grid = Grid()
        # grid.add(bar, grid_opts=opts.GridOpts(pos_bottom="60%"))
        # grid.add(line, grid_opts=opts.GridOpts(pos_top="60%"))
        # events = {
        #     "click": "function(params) { console.log(params.name); return params.name }",
        #     # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
        # }
        # yearmonth = st_pyecharts(bar, height=400, width=800, events=events)
        bar, yearmonth = print_bar(x_data, y_data, "处罚数量", "案例数量统计")
        # st.write(yearmonth)
        if yearmonth is not None:
            search_df["month"] = search_df["发文日期"].apply(lambda x: x.strftime("%Y-%m"))
            searchdfnew = search_df[search_df["month"] == yearmonth]
            # drop month column
            searchdfnew.drop(columns=["month"], inplace=True)
            # set session state
            st.session_state["search_result_csrc"] = searchdfnew

        # 图一解析开始：get max_period
        maxmonth = df_month["month"].max()
        minmonth = df_month["month"].min()
        # get total number of count
        num_total = df_month["count"].sum()
        # get total number of month count
        month_total = df_month["month"].count()
        # get average number of count per month count
        num_avg = num_total / month_total
        # format num_avg to int
        num_avg = int(num_avg)
        # get index of max count
        top1 = df_month["count"].nlargest(1)
        top1_index = df_month["count"].idxmax()
        # get month value of max count
        top1month = df_month.loc[top1_index, "month"]
        image1_text = (
            "图一解析：从"
            + minmonth
            + "至"
            + maxmonth
            + "，共发生"
            + str(num_total)
            + "起处罚事件，"
            + "平均每月发生"
            + str(num_avg)
            + "起处罚事件。其中"
            + top1month
            + "最高发生"
            + str(top1.values[0])
            + "起处罚事件"
        )
        # display total coun
        st.markdown("#####  " + image1_text)
        # yearmonthline = st_pyecharts(line, height=400, width=800, events=events)
        line, yearmonthline = print_line(x_data, sum_data, "处罚金额", "案例金额统计")
        # st.write(yearmonth)
        if yearmonthline is not None:
            search_df["month"] = search_df["发文日期"].apply(lambda x: x.strftime("%Y-%m"))
            searchdfnew = search_df[search_df["month"] == yearmonthline]
            # drop month column
            searchdfnew.drop(columns=["month"], inplace=True)
            # set session state
            st.session_state["search_result_csrc"] = searchdfnew

        # 图二解析：
        sum_data_number = 0  # 把案件金额的数组进行求和
        more_than_100 = 0  # 把案件金额大于100的数量进行统计
        case_total = 0  # 把案件的总数量进行统计

        penaltycount = df_sigle_penalty["amt"].tolist()
        for i in penaltycount:
            sum_data_number = sum_data_number + i / 10000
            if i > 100 * 10000:
                more_than_100 = more_than_100 + 1
            if i != 0:
                case_total = case_total + 1

        # for i in sum_data:
        #     sum_data_number = sum_data_number + i / 10000
        #     if i > 100 * 10000:
        #         more_than_100 = more_than_100 + 1
        # sum_data_number=round(sum_data_number,2)
        # get index of max sum
        topsum1 = df_sum["sum"].nlargest(1)
        topsum1_index = df_sum["sum"].idxmax()
        # get month value of max count
        topsum1month = df_month.loc[topsum1_index, "month"]
        image2_text = (
            "图二解析：从"
            + minmonth
            + "至"
            + maxmonth
            + "，共发生罚款案件"
            + str(case_total)
            + "起;期间共涉及处罚金额"
            + str(round(sum_data_number, 2))
            + "万元，处罚事件平均处罚金额为"
            + str(round(sum_data_number / case_total, 2))
            + "万元，其中处罚金额高于100万元处罚事件共"
            + str(more_than_100)
            + "起。"
            + topsum1month
            + "发生最高处罚金额"
            + str(round(topsum1.values[0] / 10000, 2))
            + "万元。"
        )
        st.markdown("##### " + image2_text)
    # replace blank value for column 当事人身份
    selected_peopledetail["当事人身份"].replace("", "未知", inplace=True)

    # people type count
    peopletype = (
        selected_peopledetail.groupby("当事人身份")["id"].nunique().reset_index(name="数量统计")
    )
    # sort by count
    peopletype = peopletype.sort_values(by="数量统计", ascending=False)
    # penalty type count
    penaltytype = (
        selected_peopledetail.groupby("违规类型")["id"].nunique().reset_index(name="数量统计")
    )
    # sort by count
    penaltytype = penaltytype.sort_values(by="数量统计", ascending=False)
    # law type count
    lawtype = (
        selected_lawdetail.groupby("法律法规")["id"].nunique().reset_index(name="数量统计")
    )
    # sort by count
    lawtype = lawtype.sort_values(by="数量统计", ascending=False)

    # display checkbox to show/hide graph1
    # showgraph2 = st.sidebar.checkbox("当事人身份统计", key="showgraph2")
    showgraph2 = True
    if showgraph2:
        # draw plotly bar chart
        # fig = go.Figure()
        # # make subplots of for three graphs in one figure,y label, with specs for the bar,bar,pie
        # fig = make_subplots(rows=3,
        #                     cols=1,
        #                     specs=[[{
        #                         "type": "bar"
        #                     }], [{
        #                         "type": "bar"
        #                     }], [{
        #                         "type": "bar"
        #                     }]])
        # # add trace of people type
        # fig.add_trace(go.Bar(x=peopletype['当事人身份'],
        #                      y=peopletype['数量统计'],
        #                      name='当事人身份',
        #                      legendgroup='peopletype'),
        #               row=1,
        #               col=1)
        # # add trace of penalty type
        # fig.add_trace(go.Bar(x=penaltytype['违规类型'],
        #                      y=penaltytype['数量统计'],
        #                      name='违规类型',
        #                      legendgroup='penaltytype'),
        #               row=2,
        #               col=1)
        # add trace of law type using pie chart
        # fig.add_trace(go.Pie(labels=lawtype['法律法规'],
        #                         values=lawtype['数量统计'],
        #                         textinfo='value+percent',

        #                         name='法律法规',legendgroup='lawtype'),
        #                 row=3, col=1)
        # set layout of legend
        # fig.update_layout(legend_orientation="h",
        #                     legend=dict(x=0, y=1.2))
        # fig.add_trace(go.Bar(x=lawtype['法律法规'], y=lawtype['数量统计'],
        #                      name='法律法规'),
        #               row=3,
        #               col=1)
        # # update layout of subplots
        # fig.update_layout(height=1200,
        #                   width=800,
        #                   title_text="当事人身份、违规类型、法律法规统计")
        # st.plotly_chart(fig)

        x_data1 = peopletype["当事人身份"].tolist()
        y_data1 = peopletype["数量统计"].tolist()
        # bar1 = (
        #     Bar()
        #     .add_xaxis(xaxis_data=x_data1)
        #     .add_yaxis(series_name="数量", y_axis=y_data1)
        #     .set_global_opts(
        #         xaxis_opts=opts.AxisOpts(
        #             axislabel_opts=opts.LabelOpts(is_show=True, rotate=-15)
        #         ),
        #         title_opts=opts.TitleOpts(title="当事人身份统计"),
        #     )
        # )
        # events = {
        #     "click": "function(params) { console.log(params.name); return params.name }",
        #     # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
        # }
        # display bar chart
        # peopletype_selected = st_pyecharts(bar1, width=800, height=400, events=events)
        bar1, peopletype_selected = print_bar(x_data1, y_data1, "处罚数量", "当事人身份统计")
        # get selected people type
        if peopletype_selected is not None:
            # get selected people type id
            peopletype_id = selected_peopledetail[
                selected_peopledetail["当事人身份"] == peopletype_selected
            ]["id"].unique()
            # get subsearchdf by id
            subsearchdf = search_df.loc[search_df["id"].isin(peopletype_id)]
            # set session state
            st.session_state["search_result_csrc"] = subsearchdf

        # 图三解析开始：
        dict = {"当事人身份": x_data1, "数量统计": y_data1}
        peopletype_count = pd.DataFrame(dict)  # 把人员类型的数量进行统计
        # peopletype_count.columns = ['当事人身份','数量统计']
        # pandas数据排序
        peopletype_count = peopletype_count.sort_values("数量统计", ascending=False)
        result3 = ""
        for i in range(5):
            try:
                result3 = (
                    result3
                    + str(peopletype_count.iloc[i, 0])
                    + "("
                    + str(peopletype_count.iloc[i, 1])
                    + "起),"
                )
            except Exception as e:
                print(e)
                break
        image3_text = "图三解析：处罚事件中，各当事人身份中被处罚数量排名前五分别为:" + result3
        st.markdown("##### " + image3_text)
    # showgraph3 = st.sidebar.checkbox("违规类型统计", key="showgraph3")
    showgraph3 = True
    if showgraph3:
        x_data2 = penaltytype["违规类型"].tolist()
        y_data2 = penaltytype["数量统计"].tolist()
        # bar2 = (
        #     Bar()
        #     .add_xaxis(xaxis_data=x_data2)
        #     .add_yaxis(
        #         series_name="数量",
        #         y_axis=y_data2,
        #         # label_opts=opts.LabelOpts(is_show=True, position="inside")
        #     )
        #     .set_global_opts(
        #         xaxis_opts=opts.AxisOpts(
        #             axislabel_opts=opts.LabelOpts(is_show=True, rotate=-15)
        #         ),
        #         # yaxis_opts=opts.AxisOpts(
        #         #     axislabel_opts=opts.LabelOpts(is_show=True, rotate=-15),
        #         #     # axisline_opts=opts.AxisLineOpts(
        #         #     #     linestyle_opts=opts.LineStyleOpts(color="#5793f3", )),
        #         #     # axislabel_opts=opts.LabelOpts(formatter="{value} ml"),
        #         # ),
        #         title_opts=opts.TitleOpts(title="违规类型统计"),
        #     )
        # )
        # events = {
        #     "click": "function(params) { console.log(params.name); return params.name }",
        #     # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
        # }
        # display bar chart
        # pentype_selected = st_pyecharts(bar2, width=800, height=400, events=events)
        bar2, pentype_selected = print_bar(x_data2, y_data2, "处罚数量", "违规类型统计")
        # display law type selected
        if pentype_selected is not None:
            # get unique id of selected law type
            pentype_selected_id = selected_peopledetail.loc[
                selected_peopledetail["违规类型"] == pentype_selected, "id"
            ].unique()
            # get subsearchdf by id
            subsearchdf = search_df.loc[search_df["id"].isin(pentype_selected_id)]
            # set session state
            st.session_state["search_result_csrc"] = subsearchdf

        # 图四解析开始
        penaltytype_count = penaltytype[["违规类型", "数量统计"]]  # 把违规类型的数量进行统计
        # pandas数据排序
        penaltytype_count = penaltytype_count.sort_values("数量统计", ascending=False)
        result4 = ""
        for i in range(5):
            try:
                result4 = (
                    result4
                    + str(penaltytype_count.iloc[i, 0])
                    + "("
                    + str(penaltytype_count.iloc[i, 1])
                    + "起),"
                )
            except Exception as e:
                print(e)
                break
        image4_text = "图四解析：处罚事件中，各违规类型中处罚数量排名前五分别为:" + result4[: len(result4) - 1]
        st.markdown("##### " + image4_text)

    # showgraph4 = st.sidebar.checkbox("法律法规统计", key="showgraph4")
    showgraph4 = True
    if showgraph4:
        x_data3 = lawtype["法律法规"].tolist()
        y_data3 = lawtype["数量统计"].tolist()
        # bar3 = (
        #     Bar()
        #     .add_xaxis(xaxis_data=x_data3)
        #     .add_yaxis(series_name="数量", y_axis=y_data3)
        #     .set_global_opts(
        #         xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),
        #         title_opts=opts.TitleOpts(title="法律法规统计"),
        #     )
        # )
        # events = {
        #     "click": "function(params) { console.log(params.name); return params.name }",
        #     # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
        # }
        # display bar chart
        # lawtype_selected = st_pyecharts(bar3, width=800, height=400, events=events)
        bar3, lawtype_selected = print_bar(x_data3, y_data3, "处罚数量", "法律法规统计")
        # display law type selected
        if lawtype_selected is not None:
            # get unique id of selected law type
            lawtype_selected_id = selected_lawdetail.loc[
                selected_lawdetail["法律法规"] == lawtype_selected, "id"
            ].unique()
            # get subsearchdf by id
            subsearchdf = search_df.loc[search_df["id"].isin(lawtype_selected_id)]
            # set session state
            st.session_state["search_result_csrc"] = subsearchdf

        # 图五解析开始
        lawtype_count = lawtype[["法律法规", "数量统计"]]  # 把法律法规的数量进行统计
        # pandas数据排序
        lawtype_count = lawtype_count.sort_values("数量统计", ascending=False)
        result5 = ""
        for i in range(5):
            try:
                result5 = (
                    result5
                    + str(lawtype_count.iloc[i, 0])
                    + "("
                    + str(lawtype_count.iloc[i, 1])
                    + "起),"
                )
            except Exception as e:
                print(e)
                break
        # st.markdown(
        #     "##### 图五解析:法律法规统计-不同法规维度：处罚事件中，各违规类型中处罚数量排名前五分别为:"
        #     + result5[: len(result5) - 1]
        # )
        # by具体条文
        # lawdf["数量统计"] = ""
        new_lawtype = (
            selected_lawdetail.groupby(["法律法规", "条文"])["id"]
            .nunique()
            .reset_index(name="数量统计")
        )
        # new_lawtype=lawdf.groupby(['法律法规','条文'])#%%%
        new_lawtype["法律法规明细"] = new_lawtype["法律法规"] + "(" + new_lawtype["条文"] + ")"

        lawtype_count = new_lawtype[["法律法规明细", "数量统计"]]  # 把法律法规的数量进行统计
        # pandas数据排序
        lawtype_count = lawtype_count.sort_values("数量统计", ascending=False)
        result6 = ""
        for i in range(5):
            try:
                result6 = (
                    result6
                    + str(lawtype_count.iloc[i, 0])
                    + "("
                    + str(lawtype_count.iloc[i, 1])
                    + "起),"
                )
            except Exception as e:
                print(e)
                break
        image5_text = (
            " 图五解析:法律法规统计-不同法规维度：处罚事件中，各违规类型中处罚数量排名前五分别为:"
            + result5[: len(result5) - 1]
            + "\n"
            + "法律法规统计-具体条文维度：处罚事件中，各违规类型中处罚数量排名前五分别为:"
            + result6[: len(result6) - 1]
        )
        st.markdown("##### " + image5_text)
    # display summary
    st.markdown("### 分析报告下载")

    if st.button("生成分析报告"):
        t1 = time.localtime()
        t1 = time.strftime("%Y-%m-%d %H%M%S", t1)

        image1 = bar.render(path=os.path.join(pencsrc, t1 + "image1.html"))
        image2 = line.render(path=os.path.join(pencsrc, t1 + t1 + "image2.html"))
        image3 = bar1.render(path=os.path.join(pencsrc, t1 + t1 + "image3.html"))
        image4 = bar2.render(path=os.path.join(pencsrc, t1 + t1 + "image4.html"))
        image5 = bar3.render(path=os.path.join(pencsrc, t1 + t1 + "image5.html"))
        # 做title
        title = st.session_state["keywords_csrc1"]
        title_str = ""
        # st.write(str(title[1]))
        title_str = "(分析范围：期间:" + str(title[0]) + "至" + str(title[1]) + ","
        # if len(title[0])!=0:
        #     title_str=title_str+'发文名称为:'+title[0]+'，'
        # if len(str(title[1]))!=0:
        #     title_str=title_str+'开始日期为:'+str(title[1])+'，'
        if len(str(title[2])) != 0:
            title_str = title_str + "结束日期为:" + str(title[2]) + "，"
        if len(title[3]) != 0:
            title_str = title_str + "发文单位为:" + title[3] + "，"
        if len(title[4]) != 0:
            title_str = title_str + "案件关键词为:" + title[4] + "，"
        if len(title[5]) == 2:
            #     title_str=title_str+'，'
            # else:
            title_str = title_str + "文书类型为:" + "、".join(title[5]) + "，"
        title_str = title_str[: len(title_str) - 1] + ")"
        title_str = "处罚事件分析报告\n" + title_str
        file_name = make_docx(
            title_str,
            [image1_text, image2_text, image3_text, image4_text, image5_text],
            [image1, image2, image3, image4, image5],
        )
        st.download_button(
            "下载分析报告", data=file_name.read(), file_name="分析报告.docx"
        )  # ,on_click=lambda: os.remove(file_name)
        # "下载搜索结果", data=search_dfnew.to_csv().encode("utf_8_sig"), file_name="搜索结果.csv"


def json2df(site_json):
    idls = []
    namels = []
    issueorgls = []
    filenols = []
    datels = []
    # if not empty
    if site_json["pageUtil"]["pageList"] != []:
        for i in range(20):
            idls.append(site_json["pageUtil"]["pageList"][i]["lawWritId"])
            namels.append(site_json["pageUtil"]["pageList"][i]["name"])
            issueorgls.append(site_json["pageUtil"]["pageList"][i]["issueOrgName"])
            filenols.append(site_json["pageUtil"]["pageList"][i]["fileno"])
            datels.append(site_json["pageUtil"]["pageList"][i]["dsptDate"])
    else:
        st.error("数据为空")

    eventdf = pd.DataFrame(
        {
            "id": idls,
            "name": namels,
            "issueorg": issueorgls,
            "fileno": filenols,
            "date": datels,
        }
    )
    eventdf["date"] = (
        eventdf["date"].astype(str).apply(lambda x: pd.to_datetime(x[:10], unit="s"))
    )

    return eventdf


# get sumeventdf in page number range
def get_sumeventdf(start, end):
    resultls = []
    for pageno in range(start, end + 1):
        st.info("page:" + str(pageno))
        url = BASE_URL + str(pageno)
        pp = requests.get(url, verify=False)
        ss = BeautifulSoup(pp.content, "html.parser")
        ss_json = json.loads(ss.text)
        resultdf = json2df(ss_json)
        resultls.append(resultdf)
        st.info("OK")
        time.sleep(5)

    resultsum = pd.concat(resultls).reset_index(drop=True)
    # savedf(resultsum,'sumeventdf')
    return resultsum


# get current date and time string
def get_now():
    now = datetime.datetime.now()
    now_str = now.strftime("%Y%m%d%H%M%S")
    return now_str


# update sumeventdf
def update_sumeventdf(currentsum):
    oldsum = get_csrcsum()
    if oldsum.empty:
        oldidls = []
    else:
        oldidls = oldsum["id"].tolist()
    currentidls = currentsum["id"].tolist()
    # print('oldidls:',oldidls)
    # print('currentidls:', currentidls)
    # get current idls not in oldidls
    newidls = [x for x in currentidls if x not in oldidls]
    # print('newidls:', newidls)
    # newidls=list(set(currentidls)-set(oldidls))
    newdf = currentsum[currentsum["id"].isin(newidls)]
    # if newdf is not empty, save it
    if not newdf.empty:
        newdf.reset_index(drop=True, inplace=True)
        nowstr = get_now()
        savename = "sumevent" + nowstr
        savedf(newdf, savename)
    return newdf


def title2detail(sdtitlels, detail):
    detail["文件名称"] = sdtitlels[0].text
    detail["文号"] = sdtitlels[1].text.strip()
    detail["发文日期"] = sdtitlels[2].text.strip()
    detail["文书类型"] = sdtitlels[3].text.strip()
    detail["发文单位"] = sdtitlels[4].text.strip()
    detail["原文链接"] = sdtitlels[5].text.strip()


def law2detail(sdlawls, detail):
    lawdetaills = []
    for i in range(3, len(sdlawls)):
        try:
            span = sdlawls[i]["rowspan"]
        except Exception as e:
            st.error(e)
            span = None
        if span:
            lawdetail = dict()
            itemls = []
            lawdetail["法律法规"] = sdlawls[i].text.strip()
            itemls = [sdlawls[j].text.strip() for j in range(i + 1, i + 1 + int(span))]
            lawdetail["条文"] = itemls
            lawdetaills.append(lawdetail)
        else:
            pass
    detail["处理依据"] = lawdetaills


def people2detail(sdpeoplels, detail):
    peoplenum = (len(sdpeoplels) - 6) // 5

    pdetail = dict()
    pdetaills = []
    for pno in range(peoplenum):
        #     print(pno)
        pdetail = dict()
        pdetail["当事人类型"] = sdpeoplels[6 + pno * 5].text.strip()
        pdetail["当事人名称"] = sdpeoplels[6 + pno * 5 + 1].text.strip()
        pdetail["当事人身份"] = sdpeoplels[6 + pno * 5 + 2].text.strip()
        pdetail["违规类型"] = sdpeoplels[6 + pno * 5 + 3].text.strip()
        pdetail["处罚结果"] = sdpeoplels[6 + pno * 5 + 4].text.strip()
        pdetaills.append(pdetail)
    detail["当事人信息"] = pdetaills


def fact2detail(sdfactls, detail):
    detail["案情经过"] = sdfactls[0].text.replace("\u3000", "").replace("\n", "")


# get event detail
def get_eventdetail(eventsum):
    # outputfile = pencsrc+'sdresult_0-'
    idls = eventsum["id"].tolist()

    sdresultls = []
    count = 0
    for i in idls:
        st.info("id: " + str(i))
        st.info(str(count) + " begin")
        url = urldbase + str(i)
        dd = requests.get(url, verify=False)
        sd = BeautifulSoup(dd.content, "html.parser")

        detail = dict()
        detail["id"] = i
        sdtitlels = sd.find_all(class_="text-left")

        sdlawls = sd.find_all("td", class_="text-center")

        sdpeoplels = sd.find_all(class_="table table-bordered table-condensed")[
            1
        ].find_all("td")

        sdfactls = sd.find_all(class_="pre_law")

        title2detail(sdtitlels, detail)
        law2detail(sdlawls, detail)
        people2detail(sdpeoplels, detail)
        fact2detail(sdfactls, detail)
        sdresultls.append(detail)

        # save temp result
        mod = (count + 1) % 5
        batch = (count + 1) // 5
        if mod == 0 and count > 0:
            sdresultdf = pd.DataFrame.from_dict(sdresultls)
            savename = "temp" + str(count + 1)
            savedf(sdresultdf, savename)
            st.info("batch:{} is ok".format(batch))

        st.info(str(count) + " is ok")
        count = count + 1
        time.sleep(5)

    alldf = pd.DataFrame.from_dict(sdresultls)
    # if alldf is not empty, save it
    if not alldf.empty:
        nowstr = get_now()
        savename = "sdresult" + nowstr
        savedf(alldf, savename)
    return alldf


# display event detail
def display_eventdetail(search_df):
    # get event detail
    eventdf = get_csrcdetail()
    # draw plotly figure
    display_dfmonth(search_df)
    # get search result from session
    search_dfnew = st.session_state["search_result_csrc"]
    # get searchdf idls
    idls = search_dfnew["id"].tolist()
    # get downloaddf by idls
    downloaddf = eventdf[eventdf["id"].isin(idls)]

    total = len(downloaddf)
    # st.sidebar.metric("总数:", total)
    # display search result
    st.markdown("### 搜索结果" + "(" + str(total) + "条)")
    # display download button
    st.download_button(
        "下载搜索结果",
        data=downloaddf.to_csv().encode("utf_8_sig"),
        file_name="搜索结果.csv",
    )
    # display columns
    discols = ["发文日期", "文件名称", "发文单位", "id"]
    # get display df
    display_df = downloaddf[discols]
    data = df2aggrid(display_df)
    # display data
    selected_rows = data["selected_rows"]
    if selected_rows == []:
        st.error("请先选择查看案例")
        st.stop()
    # convert selected_rows to dataframe
    selected_rows_df = pd.DataFrame(selected_rows)

    # get selected_rows_df's id
    selected_rows_id = selected_rows_df["id"].tolist()
    # get people detail
    peopledf = get_peopledetail()
    # search people detail by selected_rows_id
    selected_rows_peopledetail = peopledf[peopledf["id"].isin(selected_rows_id)]
    # display people detail
    # st.write('当事人信息')
    # st.dataframe(
    #     selected_rows_peopledetail[['当事人类型', '当事人名称', '当事人身份', '违规类型',
    #                                 '处罚结果']])
    people_data = selected_rows_peopledetail[
        ["当事人类型", "当事人名称", "当事人身份", "违规类型", "处罚结果"]
    ]
    df2echartstable(people_data, "当事人信息")
    # search event detail by selected_rows_id
    selected_rows_eventdetail = eventdf[eventdf["id"].isin(selected_rows_id)]
    # display event detail
    # st.dataframe(selected_rows_eventdetail)
    # st.table(selected_rows_eventdetail[['文件名称', '发文日期', '发文单位', '文书类型']])
    # get event detail
    event_data = selected_rows_eventdetail[
        ["发文日期", "文件名称", "发文单位", "文书类型", "案情经过", "原文链接"]
    ]
    # display event detail
    st.write("案情经过")
    # transpose
    event_data = event_data.T
    # set column name
    event_data.columns = ["案情内容"]
    st.table(event_data.astype(str))
    # df2echartstable(event_data, "案情经过")

    # get lawdf
    lawdf = get_lawdetail()
    # search lawdetail by selected_rows_id
    selected_rows_lawdetail = lawdf[lawdf["id"].isin(selected_rows_id)]
    # display lawdetail
    st.write("处罚依据")
    # st.table(selected_rows_lawdetail[['法律法规', '条文']])
    lawdata = selected_rows_lawdetail[["法律法规", "条文"]]
    lawdtl = df2aggrid(lawdata)
    selected_law = lawdtl["selected_rows"]
    if selected_law == []:
        st.error("请先选择查看监管条文")
    else:
        # get selected_law's rule name
        selected_law_name = selected_law[0]["法律法规"]
        # get selected_law's rule article
        selected_law_article = selected_law[0]["条文"]
        # get law detail by name
        ruledf = get_rulelist_byname(selected_law_name, "", "", "", "")
        # get law ids
        ids = ruledf["lawid"].tolist()
        # get law detail by id
        metadf, dtldf = get_lawdtlbyid(ids)
        # display law meta
        st.write("监管法规")
        st.table(metadf)
        # get law detail by article
        articledf = dtldf[dtldf["标题"].str.contains(selected_law_article)]
        # display law detail
        st.write("监管条文")
        st.table(articledf)


# summary of csrc
def display_summary():
    oldsum = get_csrcsum()
    # get length of old eventdf
    oldlen = len(oldsum)
    # get min and max date of old eventdf
    min_date = oldsum["date1"].min()
    max_date = oldsum["date1"].max()
    # use metric
    col1, col2 = st.columns([1, 3])
    with col1:
        st.metric("案例总数", oldlen)
    with col2:
        st.metric("案例日期范围", f"{min_date} - {max_date}")


# print bar graphs
def print_bar(x_data, y_data, y_axis_name, title):
    # draw echarts bar chart
    bar = (
        Bar()
        .add_xaxis(xaxis_data=x_data)
        .add_yaxis(series_name=y_axis_name, y_axis=y_data, yaxis_index=0)
        .set_global_opts(
            title_opts=opts.TitleOpts(title=title),
            xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),
            visualmap_opts=opts.VisualMapOpts(max_=max(y_data), min_=min(y_data)),
        )
    )
    # use events
    events = {
        "click": "function(params) { console.log(params.name); return params.name }",
        # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
    }
    # use events
    clickevent = st_pyecharts(bar, events=events, height=400)
    return bar, clickevent


# print line charts
def print_line(x_data, y_data, y_axis_name, title):
    # draw echarts line chart
    line = (
        Line()
        .add_xaxis(x_data)
        .add_yaxis(y_axis_name, y_data, label_opts=opts.LabelOpts(is_show=True))
        .set_global_opts(
            title_opts=opts.TitleOpts(title=title),
            # legend_opts=opts.LegendOpts(pos_top="48%"),
        )
    )
    # use events
    events = {
        "click": "function(params) { console.log(params.name); return params.name }",
        # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
    }
    # use events
    clickevent = st_pyecharts(line, events=events, height=400)
    return line, clickevent


def make_docx(title, text, image):  # 制作docx的函数，title以str形式传入，其他以list的形式传入，输出为字符串的形式
    document = Document()

    # st.write(title_str)
    # add title
    document.add_paragraph().add_run(title).bold = True
    # document.add_paragraph(title)
    document.styles["Normal"].font.size = Pt(12)
    document.styles["Normal"].font.name = "Times New Roman"  # 设置西文字体
    document.styles["Normal"]._element.rPr.rFonts.set(
        qn("w:eastAsia"), "FangSong"
    )  # 设置中文字体使用字
    # document.styles['Normal'].font.bold = True
    # 加粗字体

    for i, j in zip(text, image):  # [image1_text,image2_text],[image1,image2]
        # document.styles['Normal'].font.bold = False
        t = time.localtime()
        t = time.strftime("%Y-%m-%d %H%M%S", t)
        make_snapshot(driver, j, t + ".png", is_remove_html=True)  #
        document.add_paragraph(i)
        document.styles["Normal"].font.size = Pt(12)
        document.styles["Normal"].font.name = "Times New Roman"  # 设置西文字体
        document.styles["Normal"]._element.rPr.rFonts.set(
            qn("w:eastAsia"), "FangSong"
        )  # 设置中文字体使用字体2->宋体
        document.add_picture(t + ".png", width=docx.shared.Inches(5.4))  # 6英尺是最大宽度
        # print('当前图像高度', str(document.inline_shapes[0].height)+'当前图像宽度'+str(document.inline_shapes[0].width)) # 打印当前图片大小
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        os.remove(t + ".png")
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    # document.save(t1+'.docx')
    return file_stream


# get current date and time string
def get_nowdate():
    now = datetime.datetime.now()
    now_str = now.strftime("%Y%m%d")
    return now_str
